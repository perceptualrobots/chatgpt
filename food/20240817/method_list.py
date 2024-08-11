

from docx import Document

# Create a new Word document
doc = Document()

# Title of the document
doc.add_heading('Detailed Recipe Preparation Steps', 0)

# Vanilla Ice Cream with Xanthan Gum
doc.add_heading('1. Vanilla Ice Cream with Xanthan Gum', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Measure 480ml heavy cream.")
doc.add_paragraph("2. Measure 240ml whole milk.")
doc.add_paragraph("3. Measure 150g granulated sugar.")
doc.add_paragraph("4. Combine the cream, milk, and sugar in a saucepan.")
doc.add_paragraph("5. Heat the mixture over medium heat until hot but not boiling.")
doc.add_paragraph("6. Add 1 tablespoon vanilla extract to the mixture.")
doc.add_paragraph("7. Whisk in 1/2 teaspoon xanthan gum until fully combined.")
doc.add_paragraph("8. Chill the mixture in the refrigerator for several hours until cold.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. After chilling, churn the ice cream mixture in an ice cream maker according to the manufacturer's instructions.")
doc.add_paragraph("2. Once churned, transfer the ice cream to a container and freeze until firm.")

# Chocolate Avocado Truffles
doc.add_heading('2. Chocolate Avocado Truffles', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Melt 300g dark chocolate in the microwave or over a double boiler.")
doc.add_paragraph("2. Mash 1.5 ripe avocados (about 225g) until smooth.")
doc.add_paragraph("3. Combine the melted chocolate with the mashed avocado.")
doc.add_paragraph("4. Stir in 1.5 teaspoons vanilla extract.")
doc.add_paragraph("5. Chill the mixture in the refrigerator until firm.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. Once firm, roll the mixture into small truffle-sized balls.")
doc.add_paragraph("2. Roll each ball in cocoa powder to coat.")
doc.add_paragraph("3. Store the truffles in the refrigerator until ready to serve.")

# Butternut Squash with Orange Oil and Burnt Honey
doc.add_heading('3. Butternut Squash with Orange Oil and Burnt Honey', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Preheat the oven to 220°C (430°F).")
doc.add_paragraph("2. Peel, deseed, and cut 1 medium butternut squash (about 1.2kg) into wedges.")
doc.add_paragraph("3. Toss the squash with olive oil, salt, and pepper.")
doc.add_paragraph("4. Place the squash on a baking tray.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. Roast the squash in the preheated oven for 40 minutes until tender and caramelized.")
doc.add_paragraph("2. While roasting, prepare the burnt honey by heating 3 tablespoons of honey in a small saucepan until it starts to darken.")
doc.add_paragraph("3. Add 2 tablespoons of orange juice to the honey and swirl to combine.")
doc.add_paragraph("4. Drizzle the burnt honey mixture over the roasted squash before serving.")

# Lentil Bread (Flourless)
doc.add_heading('4. Lentil Bread (Flourless)', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Measure 200g red lentils and soak them in water for at least 4 hours.")
doc.add_paragraph("2. Preheat the oven to 180°C (350°F).")
doc.add_paragraph("3. Drain the soaked lentils and blend them in a food processor until smooth.")
doc.add_paragraph("4. Add 2 large eggs to the lentil mixture.")
doc.add_paragraph("5. Add 1 teaspoon baking powder and 1/2 teaspoon salt to the mixture.")
doc.add_paragraph("6. Pour the lentil mixture into a loaf tin lined with parchment paper.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. Bake the lentil bread in the preheated oven for 35-40 minutes until a toothpick inserted into the center comes out clean.")
doc.add_paragraph("2. Allow the bread to cool completely before slicing.")

# Tomato Salad
doc.add_heading('5. Tomato Salad', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Slice 800g of ripe tomatoes.")
doc.add_paragraph("2. Thinly slice 1 small red onion.")
doc.add_paragraph("3. Prepare 1/4 cup of fresh basil leaves.")

doc.add_heading('Assembly Steps:', level=2)
doc.add_paragraph("1. Arrange the sliced tomatoes on a serving plate.")
doc.add_paragraph("2. Scatter the red onion slices and basil leaves over the tomatoes.")
doc.add_paragraph("3. Dress the salad with 2 tablespoons of red wine vinegar, 3 tablespoons of olive oil, and season with salt and pepper.")

# Chicken Meatballs
doc.add_heading('6. Chicken Meatballs', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Preheat the oven to 190°C (375°F).")
doc.add_paragraph("2. In a large bowl, mix 450g ground chicken with 60g gluten-free breadcrumbs, 1 large egg, 1 teaspoon Italian seasoning, and 1 minced garlic clove.")
doc.add_paragraph("3. Season with salt and pepper to taste.")
doc.add_paragraph("4. Shape the chicken mixture into meatballs.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. Place the meatballs on a baking sheet lined with parchment paper.")
doc.add_paragraph("2. Bake the meatballs in the preheated oven for 20-25 minutes until cooked through.")

# Chorizo and Date Skewers
doc.add_heading('7. Chorizo and Date Skewers', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Preheat the oven to 200°C (400°F).")
doc.add_paragraph("2. Wrap each of 12 Medjool dates (pitted) with a slice of chorizo (200g total).")
doc.add_paragraph("3. Secure each with a toothpick.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. Arrange the skewers on a baking sheet.")
doc.add_paragraph("2. Drizzle with olive oil and roast in the preheated oven for 10-15 minutes until the chorizo is crispy.")

# Garlic Shrimp with Chili
doc.add_heading('8. Garlic Shrimp with Chili', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Slice 4 garlic cloves thinly.")
doc.add_paragraph("2. Finely chop 1 red chili.")
doc.add_paragraph("3. Chop 2 tablespoons of fresh parsley.")
doc.add_paragraph("4. Juice 1 lemon.")

doc.add_heading('Cooking/Final Steps:', level=2)
doc.add_paragraph("1. Heat 2 tablespoons of olive oil in a large pan over medium heat.")
doc.add_paragraph("2. Add the sliced garlic and chopped chili, and cook until fragrant.")
doc.add_paragraph("3. Add 450g large shrimp, peeled and deveined, and cook until pink, about 3-4 minutes.")
doc.add_paragraph("4. Stir in the chopped parsley and lemon juice, then season with salt and pepper.")
doc.add_paragraph("5. Serve hot.")

# Classic Margarita
doc.add_heading('9. Classic Margarita', level=1)
doc.add_heading('Preparation Steps:', level=2)
doc.add_paragraph("1. Measure 50ml tequila, 25ml triple sec, and 25ml fresh lime juice.")

doc.add_heading('Final Steps:', level=2)
doc.add_paragraph("1. Prepare a glass by rubbing a lime wedge around the rim and dipping it in salt to coat.")
doc.add_paragraph("2. Fill a cocktail shaker with ice cubes, add the tequila, triple sec, and lime juice, and shake well until chilled.")
doc.add_paragraph("3. Strain into the prepared glass and serve immediately.")

# Save the document
file_path_word = "/tmp/food/Recipe_Detailed_Preparation_Steps.docx"
doc.save(file_path_word)

file_path_word
