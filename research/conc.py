from openai import OpenAI
import docx
import os
from PyPDF2 import PdfReader

class ArticleSummarizer:
    def __init__(self):
        """
        Initializes the ArticleSummarizer with the OpenAI API key.

        """
        self.client = OpenAI()

    def read_file_content(self, file_path):
        """
        Reads the content of a file based on its format.

        Args:
            file_path (str): Path to the file.

        Returns:
            str: The content of the file as a string.
        """
        _, file_extension = os.path.splitext(file_path)

        if file_extension.lower() == ".txt":
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        elif file_extension.lower() == ".pdf":
            reader = PdfReader(file_path)
            return " ".join(page.extract_text() for page in reader.pages)
        elif file_extension.lower() == ".docx":
            doc = docx.Document(file_path)
            return " ".join(paragraph.text for paragraph in doc.paragraphs)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def summarize_in_chunks(self, content, chunk_size=3000):
        """
        Summarizes content in chunks if it exceeds the token limit.

        Args:
            content (str): The content to summarize.
            chunk_size (int): Approximate size of each chunk in tokens.

        Returns:
            str: Combined summary of all chunks.
        """
        chunks = [content[i:i + chunk_size] for i in range(0, len(content), chunk_size)]
        summaries = []

        for idx, chunk in enumerate(chunks):
            print(f"Processing chunk {idx + 1}/{len(chunks)}...")
            system_role = "You are an expert academic specializing in summarizing research papers."
            user_prompt = (
                f"Summarize this academic article chunk by identifying its main thesis, key arguments, "
                f"and any objections or counterarguments it addresses. Provide a clear and concise "
                f"explanation of the author's reasoning, including any philosophical concepts or theories "
                f"they rely on. If relevant, highlight the implications of the argument and how it contributes "
                f"to the broader discussion:\n\n"
                f"{chunk}"
            )

            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_role},
                    {"role": "user", "content": user_prompt}
                ]
            )
            summaries.append(response.choices[0].message.content)

        return "\n\n".join(summaries)

    def summarize_article(self, file_path):
        """
        Summarizes an academic article and provides its BibTeX citation.

        Args:
            file_path (str): Path to the academic article.

        Returns:
            str: A string containing the summary and BibTeX citation.
        """
        article_content = self.read_file_content(file_path)
        return self.summarize_in_chunks(article_content)

    def save_to_word(self, results, output_file):
        """
        Saves the summaries and BibTeX citations to a Word document.

        Args:
            results (list): A list of summaries and citations.
            output_file (str): Path to the output Word document.
        """
        doc = docx.Document()

        for idx, result in enumerate(results):
            doc.add_heading(f"Article {idx + 1}", level=1)
            doc.add_paragraph(result)

        doc.save(output_file)

    def process_articles(self, file_paths, output_file):
        """
        Processes a list of academic articles and saves the results to a Word document.

        Args:
            file_paths (list): List of file paths to the academic articles.
            output_file (str): Path to the output Word document.
        """
        results = []
        for file_path in file_paths:
            print(f"Processing: {file_path}")
            try:
                result = self.summarize_article(file_path)
                results.append(result)
            except Exception as e:
                print(f"Error processing {file_path}: {e}")

        self.save_to_word(results, output_file)


if __name__ == "__main__":
    # Example usage
    summarizer = ArticleSummarizer()

    article_paths = [
        "C:/Users/ruper/Dropbox/documents/PR-PCT/papers/consciousness/newideas/2024.12.06.627286v2.full.pdf"
    ]
    output_doc = "/tmp/summaries.docx"
    summarizer.process_articles(article_paths, output_doc)
    print(f"Summaries saved to {output_doc}")